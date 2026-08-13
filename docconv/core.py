"""docconv 核心转换逻辑。

支持的格式: doc, docx, pdf, txt, md, csv, xlsx, html
依赖: pdfplumber, python-docx, reportlab, openpyxl, html2text, markdown
docx -> pdf 优先用 LibreOffice 保排版；无 LibreOffice 时回退 reportlab 纯 Python 渲染。
.doc 读写需 LibreOffice（发行包已内置便携版，或系统安装版，exe 自动探测）。
"""
from __future__ import annotations

import csv
import html as html_lib
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import html2text
import markdown as md_lib
import pdfplumber
from docx import Document
from openpyxl import Workbook, load_workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from xml.sax.saxutils import escape as _xml_escape

SUPPORTED = {"doc", "docx", "pdf", "txt", "md", "csv", "xlsx", "html"}


def _get_cjk_font() -> str:
    """尽量注册一个支持中文的字体，失败则回退 Helvetica。"""
    candidates = [
        ("C:/Windows/Fonts/msyh.ttc", 0),
        ("C:/Windows/Fonts/simhei.ttf", 0),
        ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", 0),
        ("/System/Library/Fonts/PingFang.ttc", 0),
    ]
    for path, idx in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("CJK", path, subfontIndex=idx))
                return "CJK"
            except Exception:
                continue
    return "Helvetica"


# ------------------------- LibreOffice 探测 -------------------------
def find_libreoffice() -> "str | None":
    """查找 LibreOffice 可执行文件。

    优先级：
      1. 与 exe 同目录下的便携版 libreoffice/program/soffice.exe
         （发行包已内置，用户零配置即可用）
      2. 本地缓存目录（首次自动下载解压到的位置，见 ensure_libreoffice）
      3. 常见系统安装路径
      4. PATH 中的 soffice / libreoffice
    """
    if getattr(sys, "frozen", False):
        base = Path(sys.executable).parent
    else:
        base = Path(__file__).resolve().parent
    candidates = [base / "libreoffice" / "program" / "soffice.exe"]
    cache = _lo_cache_dir() / "program" / "soffice.exe"
    candidates.append(cache)
    candidates += [
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
    ]
    for c in candidates:
        if c.exists():
            return str(c)
    return shutil.which("soffice") or shutil.which("libreoffice")


# ------------------------- LibreOffice 自动下载 / 缓存 -------------------------
LO_VERSION = "25.8.7"
LO_MSI_URL = (
    "https://download.documentfoundation.org/libreoffice/stable/"
    f"{LO_VERSION}/win/x86_64/LibreOffice_{LO_VERSION}_Win_x86-64.msi"
)
_LO_PROGRESS_CB = None


def set_lo_progress_callback(cb) -> None:
    """设置 LO 自动下载 / 解压的进度回调，签名 cb(stage: str, percent: float|None)。"""
    global _LO_PROGRESS_CB
    _LO_PROGRESS_CB = cb


def _lo_cache_dir() -> Path:
    base = os.environ.get("LOCALAPPDATA") or tempfile.gettempdir()
    return Path(base) / "docconv" / "libreoffice"


def _report(stage: str, percent=None) -> None:
    if _LO_PROGRESS_CB:
        try:
            _LO_PROGRESS_CB(stage, percent)
        except Exception:
            pass


def ensure_libreoffice() -> "str | None":
    """确保 LibreOffice 可用：复用已存在版本，否则首次使用时自动下载并解压到缓存目录。

    返回 soffice 可执行文件路径；无网络或解压失败时返回 None。
    仅在「找不到任何 LibreOffice」时触发一次下载（约 350MB，之后缓存复用、离线可用）。
    """
    existing = find_libreoffice()
    if existing:
        return existing
    cache = _lo_cache_dir()
    bundled = cache / "program" / "soffice.exe"
    if bundled.exists():
        return str(bundled)
    cache.mkdir(parents=True, exist_ok=True)
    import urllib.request

    msi = cache / f"LibreOffice_{LO_VERSION}_Win_x86-64.msi"
    if not msi.exists():
        _report("正在下载 LibreOffice（首次使用，约 350 MB）…", 0)
        try:
            def _hook(block_num, block_size, total):
                if total:
                    _report("下载 LibreOffice…", min(100, block_num * block_size * 100 // total))
            urllib.request.urlretrieve(LO_MSI_URL, str(msi), _hook)
        except Exception as e:  # noqa: BLE001
            _report(f"LibreOffice 下载失败：{e}（可手动安装 LibreOffice 后重试）", None)
            return None
    _report("正在解压 LibreOffice（首次使用，请稍候）…", None)
    try:
        subprocess.run(
            ["msiexec", "/a", str(msi), f"TARGETDIR={cache}", "/qn", "/norestart"],
            check=True, capture_output=True,
        )
    except Exception as e:  # noqa: BLE001
        _report(f"LibreOffice 解压失败：{e}", None)
        return None
    if bundled.exists():
        _report("LibreOffice 已就绪（已缓存，后续离线可用）", 100)
        return str(bundled)
    _report("LibreOffice 解压后未找到 soffice.exe", None)
    return None


def _get_lo() -> "str | None":
    """获取（必要时自动下载）LibreOffice 路径。"""
    return find_libreoffice() or ensure_libreoffice()


def _libreoffice_convert(src: str, fmt: str, dst: str, lo: "str | None" = None) -> None:
    lo = lo or _get_lo()
    if not lo:
        raise RuntimeError(
            "未找到 LibreOffice，无法完成该转换。\n"
            "本发行版已内置便携 LibreOffice；请确认 libreoffice/ 文件夹与 exe 同级，"
            "或在 https://www.libreoffice.org/ 安装后重试。"
        )
    outdir = str(Path(dst).parent)
    cmd = [lo, "--headless", "--convert-to", fmt, "--outdir", outdir, str(src)]
    subprocess.run(cmd, check=True, capture_output=True)
    generated = Path(outdir) / (Path(src).stem + "." + fmt)
    if generated.exists() and generated.resolve() != Path(dst).resolve():
        generated.replace(dst)


# ------------------------- docx -> pdf -------------------------
def _docx_to_pdf_libreoffice(src: str, dst: str, lo: str) -> None:
    _libreoffice_convert(src, "pdf", dst, lo)


def _docx_to_pdf_reportlab(src: str, dst: str) -> None:
    """reportlab 纯 Python 渲染 docx -> pdf（无需 LibreOffice）。

    保留段落、标题层级与表格；图片不保留，复杂排版（如分栏、文本框）
    会简化为线性排版。中文依赖系统中文字体（自动探测）。
    """
    doc = Document(src)
    font = _get_cjk_font()
    ss = getSampleStyleSheet()
    normal = ParagraphStyle(
        "docconv_body", parent=ss["Normal"], fontName=font,
        fontSize=10, leading=15,
    )
    heading_styles = {}
    for i in range(1, 7):
        heading_styles[i] = ParagraphStyle(
            f"docconv_h{i}", parent=ss[f"Heading{i}"], fontName=font,
            fontSize=max(12, 20 - i), leading=max(15, 24 - i),
        )
    flow = []
    for p in doc.paragraphs:
        text = p.text
        if not text.strip():
            flow.append(Spacer(1, 6))
            continue
        style_name = (p.style.name or "") if p.style else ""
        esc = _xml_escape(text)
        if style_name == "Title":
            flow.append(Paragraph(esc, heading_styles[1]))
        elif style_name.startswith("Heading"):
            try:
                level = int(style_name[len("Heading"):])
            except ValueError:
                level = 1
            flow.append(Paragraph(esc, heading_styles[min(level, 6)]))
        else:
            flow.append(Paragraph(esc, normal))
    for table in doc.tables:
        data = [[_xml_escape(cell.text) for cell in row.cells] for row in table.rows]
        if not data:
            continue
        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("FONTNAME", (0, 0), (-1, -1), font),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        flow.append(t)
        flow.append(Spacer(1, 10))
    pdf = SimpleDocTemplate(
        str(dst), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )
    pdf.build(flow)


def docx_to_pdf(src: str, dst: str) -> None:
    """docx -> pdf：优先 LibreOffice 保排版，无 LibreOffice 时回退纯 Python。"""
    lo = find_libreoffice()
    if lo:
        _docx_to_pdf_libreoffice(src, dst, lo)
    else:
        _docx_to_pdf_reportlab(src, dst)


# ------------------------- .doc 读写 (经 LibreOffice) -------------------------
def _doc_to_temp_docx(src: str) -> str:
    """用 LibreOffice 把 .doc 转成临时 .docx，返回路径（调用方负责删除）。"""
    lo = _get_lo()
    if not lo:
        raise RuntimeError(
            "未找到 LibreOffice，无法转换 .doc 文件。\n"
            "本发行版已内置便携 LibreOffice；请确认 libreoffice/ 文件夹与 exe 同级，"
            "或在 https://www.libreoffice.org/ 安装后重试。"
        )
    tmp = tempfile.mkdtemp(prefix="docconv_")
    cmd = [lo, "--headless", "--convert-to", "docx", "--outdir", tmp, str(src)]
    subprocess.run(cmd, check=True, capture_output=True)
    return str(Path(tmp) / (Path(src).stem + ".docx"))


def doc_to_docx(src: str, dst: str) -> None:
    _libreoffice_convert(src, "docx", dst)


def doc_to_pdf(src: str, dst: str) -> None:
    _libreoffice_convert(src, "pdf", dst)


def docx_to_doc(src: str, dst: str) -> None:
    _libreoffice_convert(src, "doc", dst)


def doc_to_text(src: str, dst: str) -> None:
    d = _doc_to_temp_docx(src)
    try:
        docx_to_text(d, dst)
    finally:
        try:
            Path(d).unlink()
        except OSError:
            pass


def doc_to_markdown(src: str, dst: str) -> None:
    d = _doc_to_temp_docx(src)
    try:
        docx_to_markdown(d, dst)
    finally:
        try:
            Path(d).unlink()
        except OSError:
            pass


def doc_to_html(src: str, dst: str) -> None:
    d = _doc_to_temp_docx(src)
    try:
        docx_to_html(d, dst)
    finally:
        try:
            Path(d).unlink()
        except OSError:
            pass


# ------------------------- pdf -> docx -------------------------
def pdf_to_docx(src: str, dst: str) -> None:
    doc = Document()
    with pdfplumber.open(src) as pdf:
        for pi, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            text = page.extract_text() or ""
            if text.strip():
                for line in text.split("\n"):
                    doc.add_paragraph(line)
            for table in tables:
                if not table:
                    continue
                ncols = len(table[0]) if table else 0
                t = doc.add_table(rows=len(table), cols=ncols)
                try:
                    t.style = "Table Grid"
                except Exception:
                    pass
                for r, row in enumerate(table):
                    for c, cell in enumerate(row):
                        t.cell(r, c).text = (cell or "").strip()
            if pi < len(pdf.pages) - 1:
                doc.add_page_break()
    doc.save(dst)


# ------------------------- docx -> txt / md -------------------------
def docx_to_text(src: str, dst: str) -> None:
    doc = Document(src)
    lines = [p.text for p in doc.paragraphs]
    Path(dst).write_text("\n".join(lines), encoding="utf-8")


def docx_to_markdown(src: str, dst: str) -> None:
    doc = Document(src)
    lines = []
    for p in doc.paragraphs:
        style = (p.style.name or "") if p.style else ""
        text = p.text
        if style == "Title":
            lines.append(f"# {text}")
        elif style.startswith("Heading"):
            try:
                level = int(style[len("Heading"):])
            except ValueError:
                level = 1
            lines.append(f"{'#' * min(level, 6)} {text}")
        else:
            lines.append(text)
    Path(dst).write_text("\n".join(lines), encoding="utf-8")


# ------------------------- txt / md -> docx -------------------------
def _split_md(text: str):
    """极简 markdown 解析，返回 (kind, level, content) 列表。"""
    blocks = []
    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.lstrip().startswith("#"):
            hashes = line.lstrip().split(" ")[0]
            level = len(hashes)
            blocks.append(("heading", min(level, 6), line.lstrip("#").strip()))
        else:
            blocks.append(("p", 0, line))
    return blocks


def text_to_docx(src: str, dst: str) -> None:
    doc = Document()
    text = Path(src).read_text(encoding="utf-8")
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    doc.save(dst)


def markdown_to_docx(src: str, dst: str) -> None:
    doc = Document()
    text = Path(src).read_text(encoding="utf-8")
    for kind, level, content in _split_md(text):
        if kind == "heading":
            doc.add_heading(content, level=level)
        else:
            doc.add_paragraph(content)
    doc.save(dst)


# ------------------------- txt / md -> pdf -------------------------
def text_to_pdf(src: str, dst: str) -> None:
    font = _get_cjk_font()
    c = canvas.Canvas(str(dst), pagesize=A4)
    width, height = A4
    left = 20 * mm
    top = height - 20 * mm
    line_h = 6 * mm
    c.setFont(font, 11)
    y = top
    for line in Path(src).read_text(encoding="utf-8").split("\n"):
        if y < 20 * mm:
            c.showPage()
            y = top
            c.setFont(font, 11)
        c.drawString(left, y, line[:120])
        y -= line_h
    c.save()


def markdown_to_pdf(src: str, dst: str) -> None:
    text = Path(src).read_text(encoding="utf-8")
    cleaned = "\n".join(
        (ln.lstrip("#").strip() if ln.lstrip().startswith("#") else ln)
        for ln in text.split("\n")
    )
    fd, tmppath = tempfile.mkstemp(suffix=".txt", prefix="docconv_")
    os.close(fd)
    tmp = Path(tmppath)
    tmp.write_text(cleaned, encoding="utf-8")
    try:
        text_to_pdf(str(tmp), dst)
    finally:
        try:
            tmp.unlink()
        except OSError:
            pass


# ------------------------- 表格 csv / xlsx -------------------------
def _rows_from_csv(src: str):
    with open(src, newline="", encoding="utf-8-sig") as f:
        return [list(r) for r in csv.reader(f)]


def csv_to_xlsx(src: str, dst: str) -> None:
    wb = Workbook()
    ws = wb.active
    for row in _rows_from_csv(src):
        ws.append(row)
    wb.save(dst)


def xlsx_to_csv(src: str, dst: str) -> None:
    wb = load_workbook(src, read_only=True, data_only=True)
    ws = wb.active
    with open(dst, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            writer.writerow(["" if c is None else c for c in row])


def csv_to_docx(src: str, dst: str) -> None:
    doc = Document()
    rows = _rows_from_csv(src)
    if rows:
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        try:
            t.style = "Table Grid"
        except Exception:
            pass
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                t.cell(r, c).text = val or ""
    doc.save(dst)


def xlsx_to_docx(src: str, dst: str) -> None:
    doc = Document()
    wb = load_workbook(src, read_only=True, data_only=True)
    ws = wb.active
    rows = [list(r) for r in ws.iter_rows(values_only=True)]
    if rows:
        ncols = max((len(r) for r in rows), default=1)
        t = doc.add_table(rows=len(rows), cols=ncols)
        try:
            t.style = "Table Grid"
        except Exception:
            pass
        for r, row in enumerate(rows):
            for c in range(ncols):
                v = row[c] if c < len(row) else None
                t.cell(r, c).text = "" if v is None else str(v)
    doc.save(dst)


def _table_to_pdf(rows, dst: str) -> None:
    doc = SimpleDocTemplate(str(dst), pagesize=A4)
    ncols = max((len(r) for r in rows), default=1)
    data = [
        ["" if (c >= len(r) or r[c] is None) else str(r[c]) for c in range(ncols)]
        for r in rows
    ]
    t = Table(data)
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    doc.build([t])


def csv_to_pdf(src: str, dst: str) -> None:
    _table_to_pdf(_rows_from_csv(src), dst)


def xlsx_to_pdf(src: str, dst: str) -> None:
    wb = load_workbook(src, read_only=True, data_only=True)
    ws = wb.active
    rows = [list(r) for r in ws.iter_rows(values_only=True)]
    _table_to_pdf(rows, dst)


# ------------------------- html 互转 -------------------------
def html_to_md(src: str, dst: str) -> None:
    h = html2text.HTML2Text()
    h.body_width = 0
    raw = Path(src).read_text(encoding="utf-8")
    Path(dst).write_text(h.handle(raw), encoding="utf-8")


def html_to_text(src: str, dst: str) -> None:
    # 文本形态与 markdown 同源（html2text 默认即纯文本风格）
    html_to_md(src, dst)


def text_to_html(src: str, dst: str) -> None:
    text = Path(src).read_text(encoding="utf-8")
    esc = html_lib.escape(text)
    html = (
        "<!DOCTYPE html><html lang=\"zh\"><head><meta charset=\"utf-8\">"
        f"<title>{html_lib.escape(Path(src).stem)}</title></head>"
        f"<body><pre>{esc}</pre></body></html>"
    )
    Path(dst).write_text(html, encoding="utf-8")


def md_to_html(src: str, dst: str) -> None:
    text = Path(src).read_text(encoding="utf-8")
    body = md_lib.markdown(text, extensions=["tables", "fenced_code"])
    html = (
        "<!DOCTYPE html><html lang=\"zh\"><head><meta charset=\"utf-8\">"
        f"<title>{html_lib.escape(Path(src).stem)}</title></head>"
        f"<body>{body}</body></html>"
    )
    Path(dst).write_text(html, encoding="utf-8")


def docx_to_html(src: str, dst: str) -> None:
    doc = Document(src)
    parts = []
    for p in doc.paragraphs:
        style = (p.style.name or "") if p.style else ""
        text = html_lib.escape(p.text)
        if style == "Title":
            parts.append(f"<h1>{text}</h1>")
        elif style.startswith("Heading"):
            try:
                level = int(style[len("Heading"):])
            except ValueError:
                level = 1
            lv = min(level, 6)
            parts.append(f"<h{lv}>{text}</h{lv}>")
        else:
            parts.append(f"<p>{text}</p>")
    html = (
        "<!DOCTYPE html><html lang=\"zh\"><head><meta charset=\"utf-8\">"
        f"<title>{html_lib.escape(Path(src).stem)}</title></head>"
        f"<body>{''.join(parts)}</body></html>"
    )
    Path(dst).write_text(html, encoding="utf-8")


# ------------------------- 路由 -------------------------
_ROUTES = {
    # doc (老版 Word，经 LibreOffice；发行包内置便携版)
    ("doc", "docx"): doc_to_docx,
    ("doc", "pdf"): doc_to_pdf,
    ("doc", "txt"): doc_to_text,
    ("doc", "md"): doc_to_markdown,
    ("doc", "html"): doc_to_html,
    ("docx", "doc"): docx_to_doc,
    # 原有
    ("docx", "pdf"): docx_to_pdf,
    ("pdf", "docx"): pdf_to_docx,
    ("docx", "txt"): docx_to_text,
    ("docx", "md"): docx_to_markdown,
    ("txt", "docx"): text_to_docx,
    ("md", "docx"): markdown_to_docx,
    ("txt", "pdf"): text_to_pdf,
    ("md", "pdf"): markdown_to_pdf,
    # 表格 csv / xlsx
    ("csv", "xlsx"): csv_to_xlsx,
    ("xlsx", "csv"): xlsx_to_csv,
    ("csv", "docx"): csv_to_docx,
    ("xlsx", "docx"): xlsx_to_docx,
    ("csv", "pdf"): csv_to_pdf,
    ("xlsx", "pdf"): xlsx_to_pdf,
    # html
    ("html", "txt"): html_to_text,
    ("html", "md"): html_to_md,
    ("txt", "html"): text_to_html,
    ("md", "html"): md_to_html,
    ("docx", "html"): docx_to_html,
}


def _ext(path: str) -> str:
    return Path(path).suffix.lstrip(".").lower()


def convert(src: str, dst: str) -> str:
    """把 src 转换为 dst。两者都需带受支持的扩展名。"""
    if not os.path.exists(src):
        raise FileNotFoundError(f"源文件不存在: {src}")
    src_fmt = _ext(src)
    dst_fmt = _ext(dst)
    if src_fmt not in SUPPORTED:
        raise ValueError(f"不支持的源格式: .{src_fmt}")
    if dst_fmt not in SUPPORTED:
        raise ValueError(f"不支持的目标格式: .{dst_fmt}")
    func = _ROUTES.get((src_fmt, dst_fmt))
    if not func:
        raise NotImplementedError(f"暂不支持 {src_fmt} -> {dst_fmt}")
    func(src, dst)
    return dst
